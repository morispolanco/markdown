import streamlit as st
import json
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================================
# 1. UTILIDADES DE BAJO NIVEL (XML)
# ==========================================

def set_mirror_margins(section):
    """Activa los márgenes simétricos en el XML del documento."""
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        mirror_margins = OxmlElement('w:mirrorMargins')
        sectPr.insert(sectPr.index(cols[0]), mirror_margins)

def _insert_page_number_logic(paragraph):
    """Inserta el código XML necesario para mostrar el número de página."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)

def add_page_numbers_to_section(section):
    """Añade números de página a pares e impares, omitiendo la primera página."""
    footer_odd = section.footer
    p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
    _insert_page_number_logic(p_odd)

    footer_even = section.even_page_header
    p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
    _insert_page_number_logic(p_even)

def setup_headers(section, author, title):
    """Configura encabezados: Autor en pares, Título en impares. Times New Roman 9pt Versalita."""
    header_odd = section.header
    p_odd = header_odd.paragraphs[0] if header_odd.paragraphs else header_odd.add_paragraph()
    p_odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_odd.clear()
    run_odd = p_odd.add_run(title)
    run_odd.font.name = 'Times New Roman'
    run_odd.font.size = Pt(9)
    run_odd.font.small_caps = True

    header_even = section.even_page_header
    p_even = header_even.paragraphs[0] if header_even.paragraphs else header_even.add_paragraph()
    p_even.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_even.clear()
    run_even = p_even.add_run(author)
    run_even.font.name = 'Times New Roman'
    run_even.font.size = Pt(9)
    run_even.font.small_caps = True

def add_formatted_text(paragraph, text):
    """Maneja negritas y cursivas de Markdown."""
    if not text: return
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part: continue
        is_bold, is_italic = False, False
        clean_part = part
        if part.startswith('***') and part.endswith('***'):
            is_bold = is_italic = True
            clean_part = part[3:-3]
        elif part.startswith('**') and part.endswith('**'):
            is_bold = True
            clean_part = part[2:-2]
        elif part.startswith('*') and part.endswith('*'):
            is_italic = True
            clean_part = part[1:-1]
        run = paragraph.add_run(clean_part)
        run.bold, run.italic = is_bold, is_italic

# ==========================================
# 2. CONFIGURACIÓN EDITORIAL Y ESTILOS
# ==========================================

def setup_styles(doc):
    styles = doc.styles
    
    # Estilo Normal (Garamond)
    style_normal = styles['Normal']
    style_normal.font.name = 'Garamond'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_normal.paragraph_format.line_spacing = 1.15

    # Estilo Body Text
    if 'Body Text' not in styles: styles.add_style('Body Text', 1)
    body = styles['Body Text']
    body.font.name, body.font.size = 'Garamond', Pt(11)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Inches(0.25)
    body.paragraph_format.space_after = Pt(0)

    # Estilo First Paragraph
    if 'First Paragraph' not in styles: styles.add_style('First Paragraph', 1)
    first_p = styles['First Paragraph']
    first_p.font.name, first_p.font.size = 'Garamond', Pt(11)
    first_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    first_p.paragraph_format.first_line_indent = 0

    # Estilo Heading 1 (Capítulos principales)
    h1 = styles['Heading 1'] if 'Heading 1' in styles else styles.add_style('Heading 1', 1)
    h1.font.name = 'Aptos'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(72)
    h1.paragraph_format.space_after = Pt(36)
    h1.paragraph_format.keep_with_next = True

    # Estilo Título 2 (Heading 2)
    h2 = styles['Heading 2'] if 'Heading 2' in styles else styles.add_style('Heading 2', 1)
    h2.font.name = 'Aptos'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    return doc

def apply_layout(section, size_option, meta=None):
    if size_option == "Pocket (5.25 x 8 in)":
        section.page_width, section.page_height = Inches(5.25), Inches(8.0)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.5)
    else:
        section.page_width, section.page_height = Inches(5.5), Inches(8.5)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.6)
    
    set_mirror_margins(section)
    section.different_first_page_header_footer = True 
    
    if meta:
        setup_headers(section, meta['author'], meta['title'])

# ==========================================
# 3. MOTOR DE CONVERSIÓN
# ==========================================

def run_book_conversion(md_text, meta, size_option):
    doc = Document()
    setup_styles(doc)
    doc.settings.odd_and_even_pages_header_footer = True
    
    # --- PÁGINA 1: PORTADA ---
    section = doc.sections[0]
    apply_layout(section, size_option)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(120)
    run_t = p_title.add_run(meta['title'].upper())
    run_t.font.size = Pt(28)
    run_t.bold = True

    if meta.get('subtitle'):
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(6)
        run_s = p_sub.add_run(meta['subtitle'])
        run_s.font.size = Pt(14)
        run_s.italic = True

    p_author = doc.add_paragraph()
    p_author.paragraph_format.space_before = Inches(2.0)
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = p_author.add_run(meta['author'])
    run_a.font.size = Pt(16)

    if meta.get('publisher'):
        p_pub = doc.add_paragraph()
        p_pub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pub.paragraph_format.space_before = Pt(6)
        run_p = p_pub.add_run(meta['publisher'])
        run_p.font.size = Pt(12)
        run_p.italic = True

    # --- PÁGINA 2: COPYRIGHT ---
    doc.add_page_break()
    p_copy = doc.add_paragraph()
    p_copy.paragraph_format.space_before = Inches(4.5)
    
    if meta.get('copyright'):
        copyright_text = meta['copyright']
    else:
        copyright_text = (
            f"{meta['title']}\n"
            f"Copyright © {meta['year']} {meta['author']}\n"
            "All rights reserved.\n\n"
            f"Editorial: {meta.get('publisher', 'N/A')}\n"
            f"ISBN: {meta.get('isbn', '________________')}"
        )
    
    run_c = p_copy.add_run(copyright_text)
    run_c.font.size = Pt(9)
    
    # --- SECCIÓN: SOBRE EL AUTOR (Impar) ---
    section_bio = doc.add_section(WD_SECTION_START.ODD_PAGE)
    apply_layout(section_bio, size_option)
    section_bio.different_first_page_header_footer = True
    
    p_bio_title = doc.add_paragraph(style='Heading 1')
    add_formatted_text(p_bio_title, "Sobre el autor")
    
    p_bio_content = doc.add_paragraph(style='First Paragraph')
    bio_text = meta.get('about_author', "")
    if not bio_text: bio_text = "Biografía no proporcionada."
    add_formatted_text(p_bio_content, bio_text)

    # --- SECCIÓN: ÍNDICE (Impar) ---
    section_index = doc.add_section(WD_SECTION_START.ODD_PAGE)
    apply_layout(section_index, size_option)
    section_index.different_first_page_header_footer = True
    
    p_idx = doc.add_paragraph()
    p_idx.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_idx.paragraph_format.space_before = Pt(72)
    run_idx = p_idx.add_run("ÍNDICE")
    run_idx.font.size = Pt(18)
    run_idx.bold = True
    
    # --- CUERPO ---
    lines = md_text.split('\n')
    is_after_heading = False
    
    for line in lines:
        clean = line.strip()
        if not clean: continue
        
        header_match = re.match(r'^(#+)\s*(.*)$', clean)
        if header_match:
            level = len(header_match.group(1))
            title_text = header_match.group(2)
            if level == 1:
                new_sect = doc.add_section(WD_SECTION_START.ODD_PAGE)
                apply_layout(new_sect, size_option, meta)
                add_page_numbers_to_section(new_sect)
                new_sect.different_first_page_header_footer = True 
                p = doc.add_paragraph(style='Heading 1')
                add_formatted_text(p, title_text.upper())
            else:
                p = doc.add_paragraph(style=f'Heading {min(level, 5)}')
                add_formatted_text(p, title_text)
            is_after_heading = True
        else:
            style = 'First Paragraph' if is_after_heading else 'Body Text'
            p = doc.add_paragraph(style=style)
            add_formatted_text(p, clean)
            is_after_heading = False

    bio_stream = BytesIO()
    doc.save(bio_stream)
    return bio_stream.getvalue()

# ==========================================
# 4. INTERFAZ STREAMLIT
# ==========================================

st.set_page_config(page_title="Maquetador Editorial Pro", layout="centered")

# Inicialización robusta de session_state
defaults = {
    'title': "Mi Gran Novela",
    'subtitle': "",
    'author': "Nombre del Autor",
    'publisher': "Editorial Desconocida",
    'year': "2025",
    'isbn': "",
    'copyright': "",
    'about_author': "Biografía del autor...",
    'manuscript': ""
}

if 'book_data' not in st.session_state:
    st.session_state.book_data = defaults
else:
    for key, val in defaults.items():
        if key not in st.session_state.book_data:
            st.session_state.book_data[key] = val

st.title("📖 Maquetador Editorial Profesional")

with st.sidebar:
    st.header("📂 Importación")
    up_json = st.file_uploader("1. Ficha JSON", type=["json"])
    
    if up_json:
        try:
            raw_data = up_json.read().decode("utf-8")
            data = json.loads(raw_data)
            
            if not isinstance(data, dict):
                st.error("El JSON debe ser un objeto { ... }")
            else:
                # Mapeo de claves incluyendo la nueva "authorBio"
                mapping = {
                    'titulo': 'title', 'title': 'title', 
                    'subtitulo': 'subtitle', 'subtitle': 'subtitle',
                    'autor': 'author', 'author': 'author', 
                    'editorial': 'publisher', 'publisher': 'publisher',
                    'año': 'year', 'year': 'year', 
                    'isbn': 'isbn', 
                    'copyright': 'copyright',
                    'biografia': 'about_author', 'biografía': 'about_author', 
                    'bio': 'about_author', 'sobre_autor': 'about_author', 
                    'about_author': 'about_author', 'authorbio': 'about_author'
                }
                
                for k, v in data.items():
                    key_low = k.lower()
                    if key_low in mapping:
                        st.session_state.book_data[mapping[key_low]] = str(v) if v is not None else ""
                
                st.success("¡Datos cargados correctamente!")
                st.rerun()
        except json.JSONDecodeError:
            st.error("Error de sintaxis en el archivo JSON.")
        except Exception as e:
            st.error(f"Error inesperado: {e}")

    st.divider()
    up_md = st.file_uploader("2. Manuscrito (.md, .txt)", type=["md", "txt"])
    if up_md:
        try:
            content = up_md.read().decode("utf-8")
            st.session_state.book_data['manuscript'] = content
            st.success("Texto cargado")
        except: st.error("Error al leer manuscrito")

with st.expander("📝 Metadatos", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        m_title = st.text_input("Título", value=st.session_state.book_data['title'])
        m_author = st.text_input("Autor", value=st.session_state.book_data['author'])
        m_pub = st.text_input("Editorial", value=st.session_state.book_data['publisher'])
    with col2:
        m_year = st.text_input("Año", value=st.session_state.book_data['year'])
        size_mode = st.selectbox("Formato", ["Pocket (5.25 x 8 in)", "Trade (5.5 x 8.5 in)"])
        m_isbn = st.text_input("ISBN", value=st.session_state.book_data['isbn'])
    
    m_sub = st.text_input("Subtítulo", value=st.session_state.book_data['subtitle'])
    m_copyright = st.text_area("Copyright / Créditos Legales", value=st.session_state.book_data['copyright'], height=80)
    m_about = st.text_area("Sobre el autor (Biografía)", value=st.session_state.book_data['about_author'], height=120)

st.subheader("🖋️ Manuscrito")
editor_content = st.text_area("Contenido del Manuscrito", value=st.session_state.book_data.get('manuscript', ""), height=300)

if st.button("🚀 Generar Libro", use_container_width=True):
    if editor_content and m_title:
        bundle = {
            'title': m_title, 'subtitle': m_sub, 'author': m_author, 
            'publisher': m_pub, 'year': m_year, 'isbn': m_isbn, 
            'copyright': m_copyright, 'about_author': m_about
        }
        docx_bytes = run_book_conversion(editor_content, bundle, size_mode)
        st.success("¡Libro generado!")
        st.download_button(f"📥 Descargar {m_title}.docx", docx_bytes, f"{m_title.replace(' ', '_')}.docx")
